"""
多格式文档加载器

支持 PDF、Word (.docx)、纯文本 (.txt)、Markdown (.md) 文件的加载，
统一输出为 LangChain Document 对象列表。
"""

from pathlib import Path

from langchain_core.documents import Document


class DocumentLoader:
    """统一文档加载入口，根据文件扩展名自动选择加载器。"""

    # 支持的文件格式及其 MIME 描述
    SUPPORTED_FORMATS: dict[str, str] = {
        ".pdf": "PDF 文档",
        ".docx": "Word 文档",
        ".txt": "纯文本",
        ".md": "Markdown",
    }

    @classmethod
    def is_supported(cls, file_path: str | Path) -> bool:
        """检查文件格式是否受支持。"""
        return Path(file_path).suffix.lower() in cls.SUPPORTED_FORMATS

    @classmethod
    def load(cls, file_path: str | Path) -> list[Document]:
        """
        加载单个文件，返回 Document 列表。

        Args:
            file_path: 文件路径

        Returns:
            包含文件内容的 Document 列表（PDF 按页拆分）

        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 不支持的文件格式
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {path}")

        suffix = path.suffix.lower()
        loader_map = {
            ".pdf": cls._load_pdf,
            ".docx": cls._load_docx,
            ".txt": cls._load_text,
            ".md": cls._load_text,
        }

        handler = loader_map.get(suffix)
        if handler is None:
            supported = ", ".join(cls.SUPPORTED_FORMATS.keys())
            raise ValueError(f"不支持的格式 '{suffix}'，目前支持: {supported}")

        docs = handler(path)
        # 统一注入来源元数据
        for doc in docs:
            doc.metadata.setdefault("source", str(path))
            doc.metadata.setdefault("file_type", suffix)
        return docs

    @classmethod
    def load_dir(cls, dir_path: str | Path, recursive: bool = True) -> list[Document]:
        """
        批量加载目录下的所有支持格式文件。

        Args:
            dir_path: 目录路径
            recursive: 是否递归子目录

        Returns:
            所有文件的 Document 合并列表
        """
        directory = Path(dir_path)
        if not directory.is_dir():
            raise NotADirectoryError(f"目录不存在: {directory}")

        pattern = "**/*" if recursive else "*"
        all_docs: list[Document] = []

        for file_path in sorted(directory.glob(pattern)):
            if file_path.is_file() and cls.is_supported(file_path):
                all_docs.extend(cls.load(file_path))

        return all_docs

    # ── 私有加载方法 ─────────────────────────────────────────

    @staticmethod
    def _load_pdf(path: Path) -> list[Document]:
        """使用 PyMuPDF 加载 PDF，每页生成一个 Document。"""
        import pymupdf

        docs: list[Document] = []
        pdf = pymupdf.open(str(path))
        for page_num in range(len(pdf)):
            page = pdf[page_num]
            text = page.get_text("text")
            if text.strip():
                docs.append(
                    Document(
                        page_content=text,
                        metadata={"page": page_num + 1, "total_pages": len(pdf)},
                    )
                )
        pdf.close()
        return docs

    @staticmethod
    def _load_docx(path: Path) -> list[Document]:
        """使用 python-docx 加载 Word 文档。"""
        from docx import Document as DocxDocument

        docx = DocxDocument(str(path))
        paragraphs = [p.text for p in docx.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)

        if not full_text.strip():
            return []

        return [Document(page_content=full_text, metadata={})]

    @staticmethod
    def _load_text(path: Path) -> list[Document]:
        """加载纯文本 / Markdown 文件。"""
        text = path.read_text(encoding="utf-8")
        if not text.strip():
            return []
        return [Document(page_content=text, metadata={})]
